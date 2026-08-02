import re
from pathlib import Path


def _sections(body):
    parts = re.split(r"\n(?=##?\s)", body)
    out = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        m = re.match(r"#+\s*(.+)\n?(.*)", p, re.S)
        if m:
            out.append((m.group(1).strip(), m.group(2).strip()))
        else:
            out.append(("", p))
    return out


def _latin(s):
    return s.encode("latin-1", "ignore").decode()


def write_md(path, title, body):
    Path(path).write_text(f"# {title}\n\n{body}\n", encoding="utf-8")


def write_docx(path, title, body):
    from docx import Document
    d = Document()
    d.add_heading(title, 0)
    for head, text in _sections(body):
        if head:
            d.add_heading(head, 1)
        if text:
            d.add_paragraph(text)
    d.save(str(path))


def write_pdf(path, title, body):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(True, 18)
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, _latin(title), ln=True)
    pdf.ln(4)
    for head, text in _sections(body):
        if head:
            pdf.set_font("Helvetica", "B", 13)
            pdf.cell(0, 9, _latin(head), ln=True)
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 6, _latin(text or ""))
        pdf.ln(3)
    pdf.output(str(path))